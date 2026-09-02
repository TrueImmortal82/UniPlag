import codecs
import unicodedata
import re
from pathlib import Path


def extract_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    raw = path.read_bytes()

    if not raw:
        return ""

    # Check for RTF signature
    if raw.startswith(b"{\\rtf"):
        text = _from_rtf(raw)
    elif suffix == ".pdf":
        text = _from_pdf(path, raw)
    elif suffix in {".docx", ".doc"}:
        text = _from_docx_or_doc(path, raw)
    elif suffix == ".odt":
        text = _from_odt(path, raw)
    elif suffix in {".txt", ".md", ".rtf", ".py", ".c", ".cpp", ".java", ".js", ".html", ".htm", ".json", ".csv", ".log", ""}:
        text = _decode_txt(raw)
    else:
        # Universal text fallback
        try:
            text = _decode_txt(raw)
        except Exception:
            raise ValueError(f"Неподдерживаемый формат: {suffix}")

    # Aris Directive (v0.4.1, FIX): нормализуем юникод на выходе —
    # NFD-фрагменты превращаются в NFC, чтобы не ломать шинглы/ICG.
    return unicodedata.normalize("NFC", text or "")


# charset_normalizer верно детектит utf-16/utf-32 (в т.ч. без BOM) и utf-8-sig,
# но для русских legacy-кодировок (koi8-r, cp866, iso-8859-5) он систематически
# ошибается на реальных сэмплах. Поэтому авто-детект доверяем ТОЛЬКО для utf-family,
# а кириллический legacy берём строгим cp1251 (частый случай для RU-доков).
_AUTODETECT_SAFE = {"utf-8", "utf-8-sig", "utf-16", "utf-16-le", "utf-16-be",
                    "utf-32", "utf-32-le", "utf-32-be", "utf16", "utf16le", "utf16be"}


def _norm_enc(name: str) -> str:
    return name.lower().strip().replace("_", "-")


def _has_garbage_controls(text: str, threshold: float = 0.03) -> bool:
    """BOM-less UTF-16 (кириллица) обычно валиден как UTF-8: младшие байты — ASCII-
    контрольные символы (0x04, 0x1C, ...). Если в "utf-8" распозналось слишком много
    control-кодов — это артефакт, а не текст; уходим в авто-детект."""
    if not text:
        return False
    n = sum(1 for ch in text if ord(ch) < 0x20 and ch not in "\t\n\r\x0b\x0c")
    return n / len(text) > threshold


def _decode_txt(raw: bytes) -> str:
    """UTF-8 -> charset_normalizer (utf-family) -> cp1251 -> cp866 -> koi8-r -> NFC."""
    if raw.startswith(codecs.BOM_UTF8):
        return unicodedata.normalize("NFC", raw[len(codecs.BOM_UTF8):].decode("utf-8", errors="ignore"))
    try:
        text = raw.decode("utf-8")
        if not _has_garbage_controls(text):
            return unicodedata.normalize("NFC", text)
    except UnicodeDecodeError:
        pass

    try:
        from charset_normalizer import from_bytes
        match = from_bytes(raw).best()
        if match is not None and match.encoding is not None and _norm_enc(match.encoding) in _AUTODETECT_SAFE:
            text = raw.decode(match.encoding, errors="ignore")
            if text.startswith("\ufeff"):
                text = text[1:]
            return unicodedata.normalize("NFC", text)
    except (LookupError, UnicodeDecodeError, ImportError):
        pass

    try:
        return unicodedata.normalize("NFC", raw.decode("cp1251"))
    except UnicodeDecodeError:
        try:
            return unicodedata.normalize("NFC", raw.decode("cp866"))
        except UnicodeDecodeError:
            return unicodedata.normalize("NFC", raw.decode("utf-8", errors="replace"))


def _from_txt(path: Path) -> str:
    return _decode_txt(path.read_bytes())


def _from_pdf(path: Path, raw: bytes) -> str:
    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz
        except ImportError as e:
            raise RuntimeError("Для PDF установите pymupdf: pip install pymupdf") from e
    
    try:
        with fitz.open(stream=raw, filetype="pdf") as pdf:
            pages = [page.get_text() for page in pdf]
            text = "\n".join(p for p in pages if p.strip())
            if text.strip():
                return text
    except Exception:
        pass

    return _extract_strings_from_binary(raw)


def _from_docx_or_doc(path: Path, raw: bytes) -> str:
    try:
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("\t".join(cells))
        if parts:
            return "\n".join(parts)
    except Exception:
        pass

    # Fallback for legacy binary .doc or corrupted docx
    return _extract_strings_from_binary(raw)


def _from_odt(path: Path, raw: bytes) -> str:
    import zipfile
    from xml.etree import ElementTree as ET

    try:
        with zipfile.ZipFile(path) as z:
            root = ET.fromstring(z.read("content.xml"))
        paras = []
        for p in root.iter("{urn:oasis:names:tc:opendocument:xmlns:text:1.0}p"):
            paras.append("".join(p.itertext()))
        return "\n".join(re.sub(r"\s+", " ", p) for p in paras)
    except Exception:
        return _extract_strings_from_binary(raw)


def _from_rtf(raw: bytes) -> str:
    """Extracts plain text from Rich Text Format (RTF) stream."""
    try:
        text = raw.decode("latin1", errors="ignore")
        # Remove RTF meta groups
        pattern = r"\{\*?\\(?:fonttbl|colortbl|stylesheet|info|pict|header|footer)[^}]*\}"
        text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.DOTALL)
        
        # Replace hex escapes \'hh (e.g. Russian in cp1251)
        def replace_hex(m):
            try:
                return bytes.fromhex(m.group(1)).decode("cp1251", errors="ignore")
            except Exception:
                return ""
        text = re.sub(r"\\\'([0-9a-fA-F]{2})", replace_hex, text)

        # Replace Unicode escapes \uN
        def replace_uni(m):
            code = int(m.group(1))
            if code < 0:
                code += 65536
            return chr(code)
        text = re.sub(r"\\u(-?\d+)\??", replace_uni, text)

        # Remove control words and brackets
        text = re.sub(r"\\[a-zA-Z]+(?:\s+|-?\d+)?|\{|\}", " ", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n", "\n\n", text)
        return text.strip()
    except Exception:
        return _decode_txt(raw)


def _extract_strings_from_binary(raw: bytes) -> str:
    """Extracts printable Unicode (UTF-16LE) and Cyrillic/ASCII strings from binary streams."""
    # 1. UTF-16LE strings (Word 97-2003 .doc format)
    utf16_runs = []
    curr = bytearray()
    for i in range(0, len(raw) - 1, 2):
        chunk = raw[i:i+2]
        try:
            ch = chunk.decode("utf-16le")
            if ch.isprintable() or ch in "\n\r\t ":
                curr.extend(chunk)
            else:
                if len(curr) >= 8:
                    utf16_runs.append(curr.decode("utf-16le", errors="ignore"))
                curr = bytearray()
        except Exception:
            curr = bytearray()
    if curr and len(curr) >= 8:
        utf16_runs.append(curr.decode("utf-16le", errors="ignore"))

    utf16_text = " ".join(utf16_runs)

    # 2. CP1251 / ASCII strings
    ascii_runs = re.findall(rb"[\x20-\x7E\xC0-\xFF\xA8\xB8]{4,}", raw)
    ascii_list = []
    for r in ascii_runs:
        try:
            ascii_list.append(r.decode("cp1251", errors="ignore"))
        except Exception:
            pass
    ascii_text = " ".join(ascii_list)

    result = utf16_text if len(utf16_text) > len(ascii_text) else ascii_text
    result = re.sub(r"[ \t]+", " ", result)
    return result.strip()
